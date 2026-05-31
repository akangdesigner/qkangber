# -*- coding: utf-8 -*-
"""Build a Word (.docx) submission file from the submission HTML.

- Embeds the real local images (blog-drafts/01-vibe-coding-security-submission/images)
- Preserves <strong> (bold) and red emphasis spans (#c0392b)
- Replaces the bottom custom author bio with the standard "公版" bio from AuthorCard.tsx
"""
import os
import io
from html.parser import HTMLParser
from PIL import Image as PILImage
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = os.path.dirname(os.path.abspath(__file__))
SUB_DIR = os.path.join(BASE, "01-vibe-coding-security-submission")
HTML_PATH = os.path.join(SUB_DIR, "01-vibe-coding-security-submission.html")
IMG_DIR = os.path.join(SUB_DIR, "images")
OUT_PATH = os.path.join(SUB_DIR, "01-vibe-coding-security-submission.docx")

RED = RGBColor(0xC0, 0x39, 0x2A)

# ImgBB URL -> local file (from _imgbb-urls.json mapping)
URL_TO_LOCAL = {
    "https://i.ibb.co/LfhBWww/dc8a40e9efd4.jpg": "five-doors-sub.jpg",
    "https://i.ibb.co/LDL6xSF3/3afb63075448.jpg": "fixes-map-sub.jpg",
    "https://i.ibb.co/yn16wLS6/064659d56e59.jpg": "claude-scan-sub.jpg",
}

# 公版 author bio (from components/shared/AuthorCard.tsx)
AUTHOR_BIO = (
    "Q kangber｜商業自動化碩士，現職行銷公司 AI 流程開發工程師，同時協助企業 AI "
    "轉型與內部培訓，並於職涯平台擔任培訓講師及舉辦實體講座。不相信 AI 能取代你的"
    "判斷，但相信它能讓你的想法更快落地——我在做的，是找到人與 AI 之間那個最有效率的分工點。"
)


class Block:
    def __init__(self, kind):
        self.kind = kind          # h1/h2/h3/p/figure/figcaption
        self.runs = []            # list of (text, bold, red)
        self.img = None           # local image filename
        self.alt = ""


class DocParser(HTMLParser):
    def __init__(self):
        super().__init__()
        self.blocks = []
        self.cur = None
        self.bold = 0
        self.red = 0

    def handle_starttag(self, tag, attrs):
        a = dict(attrs)
        if tag in ("h1", "h2", "h3", "p", "figcaption"):
            self.cur = Block(tag)
            self.blocks.append(self.cur)
        elif tag == "strong":
            self.bold += 1
        elif tag == "span":
            style = a.get("style", "")
            if "c0392b" in style.lower():
                self.red += 1
            else:
                self.red += 0
                # track non-red span so we can balance in endtag
                self._nonred = getattr(self, "_nonred", 0) + 1
        elif tag == "img":
            src = a.get("src", "")
            local = URL_TO_LOCAL.get(src)
            b = Block("figure")
            b.img = local
            b.alt = a.get("alt", "")
            self.blocks.append(b)
            self.cur = None

    def handle_endtag(self, tag):
        if tag == "strong":
            self.bold = max(0, self.bold - 1)
        elif tag == "span":
            # red spans were the only ones we incremented; non-red handled separately
            if self.red > 0 and getattr(self, "_lastspan_red", True):
                pass
            self.red = max(0, self.red - 1) if self.red > 0 else 0
        elif tag in ("h1", "h2", "h3", "p", "figcaption"):
            self.cur = None

    def handle_data(self, data):
        text = data.replace("\n", " ")
        # collapse runs of whitespace
        while "  " in text:
            text = text.replace("  ", " ")
        if not text.strip() and not (self.cur and self.cur.runs):
            return
        if self.cur is None:
            return
        self.cur.runs.append((text, self.bold > 0, self.red > 0))


with open(HTML_PATH, encoding="utf-8") as f:
    html = f.read()

# isolate body
body = html.split("<body>", 1)[1].split("</body>", 1)[0]

p = DocParser()
p.feed(body)

doc = Document()
# base font for CJK
style = doc.styles["Normal"]
style.font.name = "Microsoft JhengHei"
style.font.size = Pt(11)
# ensure east-asian font applied
from docx.oxml.ns import qn
style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")


def add_runs(par, runs):
    for text, bold, red in runs:
        if text == "":
            continue
        r = par.add_run(text)
        r.bold = bold
        if red:
            r.font.color.rgb = RED


def is_author_block(block):
    if block.kind != "p":
        return False
    txt = "".join(t for t, _, _ in block.runs)
    return txt.strip().startswith("關於作者")


author_done = False
for b in p.blocks:
    if b.kind == "figure":
        if b.img:
            path = os.path.join(IMG_DIR, b.img)
            # Re-encode via Pillow so the JPEG carries a JFIF header that
            # python-docx can recognise (originals start with SOI+DQT only).
            with PILImage.open(path) as im:
                im = im.convert("RGB")
                buf = io.BytesIO()
                im.save(buf, format="JPEG", quality=90)
                buf.seek(0)
            par = doc.add_paragraph()
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = par.add_run()
            run.add_picture(buf, width=Inches(6.0))
        continue

    text = "".join(t for t, _, _ in b.runs).strip()
    if not text:
        continue

    if b.kind == "h1":
        h = doc.add_heading(level=0)
        add_runs(h, b.runs)
    elif b.kind == "h2":
        doc.add_heading(text, level=1)
    elif b.kind == "h3":
        doc.add_heading(text, level=2)
    elif b.kind == "figcaption":
        par = doc.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = par.add_run(text)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    elif b.kind == "p":
        if is_author_block(b):
            # replace with 公版 author info
            doc.add_heading("關於作者", level=1)
            doc.add_paragraph(AUTHOR_BIO)
            author_done = True
        else:
            par = doc.add_paragraph()
            add_runs(par, b.runs)

doc.save(OUT_PATH)
print("Saved:", OUT_PATH)
print("Author bio replaced with 公版:", author_done)
